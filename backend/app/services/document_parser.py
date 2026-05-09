"""
Unified Document Parser

Supports:
  .md    — Deep Markdown (headers, code, links, sections, metadata)
  .txt   — Plain text with smart paragraph detection
  .pdf   — pdfplumber: text + tables + layout analysis
  .docx  — python-docx: paragraphs + headings + tables + comments
  .py    — AST: functions, classes, imports, docstrings
  .js/.ts — Regex + heuristics: functions, imports, JSDoc
  .json  — Schema analysis + value extraction
  .csv   — Pandas: columns, stats, entity detection
  .html  — BeautifulSoup: main content, strip nav/footer

Each parser returns a standard ParsedDocument so the rest of
the pipeline (entity extraction, graph building, vectorization)
doesn't need to know which format was used.
"""

import ast
import csv
import json
import logging
import re
import zipfile
from dataclasses import asdict, dataclass, field
from xml.etree import ElementTree
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)


# Parsed document shape

@dataclass
class Section:
    title: str
    level: int          # heading depth: 1 = H1, 0 = no heading
    content: str


@dataclass
class CodeBlock:
    language: str
    code: str
    context: str = ""   # surrounding section title


@dataclass
class TableData:
    headers: list[str]
    rows: list[list[str]]
    caption: str = ""


@dataclass
class ParsedDocument:
    # Identity
    title: str
    file_path: str
    format: str          # md, txt, pdf, docx, py, js, ts, json, csv, html

    # Content
    raw_text: str        # full text for LLM extraction
    sections: list[Section] = field(default_factory=list)
    chunks: list[dict]   = field(default_factory=list)   # [{text, type, meta}]
    code_blocks: list[CodeBlock] = field(default_factory=list)
    tables: list[TableData] = field(default_factory=list)
    links: list[dict]    = field(default_factory=list)   # [{text, url}]
    images: list[dict]   = field(default_factory=list)   # [{alt, url}]
    list_items: list[dict] = field(default_factory=list)  # [{text, level, ordered}]

    # Structural elements (format-specific)
    imports: list[str]   = field(default_factory=list)   # code imports
    functions: list[str] = field(default_factory=list)   # function names
    classes: list[str]   = field(default_factory=list)   # class names
    entities: list[dict] = field(default_factory=list)   # lightweight extracted entities

    # Metadata
    metadata: dict       = field(default_factory=dict)
    word_count: int = 0
    reading_time_min: int = 0


# Shared chunking

def _chunk(text: str, chunk_size: int = 800, overlap: int = 150,
           chunk_type: str = "paragraph", meta: dict = None) -> list[dict]:
    """Split text into overlapping chunks for vector embedding."""
    meta = meta or {}
    if not text.strip():
        return []
    if overlap >= chunk_size:
        overlap = max(0, chunk_size // 5)
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        piece = text[start:end].strip()
        if piece:
            chunks.append({"text": piece, "type": chunk_type,
                           "start": start, "end": min(end, len(text)), **meta})
        start = end - overlap
    return chunks


# Markdown files

class MarkdownParser:
    """
    Deep Markdown parser.
    Extracts structure, code, links, frontmatter, and builds semantic chunks.
    """

    def parse(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8")

        # Frontmatter (YAML between ---)
        frontmatter = {}
        fm_match = re.match(r'^---\n(.*?)\n---\n', text, re.DOTALL)
        if fm_match:
            try:
                import yaml
                frontmatter = yaml.safe_load(fm_match.group(1)) or {}
            except ImportError:
                log.debug("PyYAML not installed; ignoring Markdown frontmatter")
            except yaml.YAMLError as exc:
                log.warning("Invalid Markdown frontmatter in %s: %s", path.name, exc)
            text = text[fm_match.end():]

        sections = self._extract_sections(text)
        code_blocks = self._extract_code_blocks(text)
        links = self._extract_links(text)
        images = self._extract_images(text)
        list_items = self._extract_list_items(text)
        clean_text = self._strip_code_blocks(text)
        chunks = self._build_chunks(sections, code_blocks)

        title = (frontmatter.get("title")
                 or next((s.title for s in sections if s.level == 1), path.stem))

        words = len(clean_text.split())
        return ParsedDocument(
            title        = title,
            file_path    = str(path),
            format       = "md",
            raw_text     = clean_text,
            sections     = sections,
            chunks       = chunks,
            code_blocks  = code_blocks,
            links        = links,
            images       = images,
            list_items   = list_items,
            metadata     = {
                **frontmatter,
                "filename": path.name,
                "section_count": len(sections),
                "list_count": len(list_items),
                "image_count": len(images),
            },
            word_count   = words,
            reading_time_min = max(1, words // 250),
        )

    def _extract_sections(self, text: str) -> list[Section]:
        sections, buf, current_title, current_level = [], [], "Introduction", 0
        for line in text.splitlines():
            m = re.match(r'^(#{1,6})\s+(.+)$', line)
            if m:
                if buf:
                    sections.append(Section(current_title, current_level, "\n".join(buf).strip()))
                    buf = []
                current_level = len(m.group(1))
                current_title = m.group(2).strip()
            else:
                buf.append(line)
        if buf:
            sections.append(Section(current_title, current_level, "\n".join(buf).strip()))
        return sections

    def _extract_code_blocks(self, text: str) -> list[CodeBlock]:
        blocks = []
        for m in re.finditer(r'```(\w*)\n(.*?)```', text, re.DOTALL):
            lang = m.group(1) or "text"
            code = m.group(2).strip()
            if code:
                blocks.append(CodeBlock(language=lang, code=code))
        return blocks

    def _extract_links(self, text: str) -> list[dict]:
        return [
            {"text": m.group(1), "url": m.group(2)}
            for m in re.finditer(r'(?<!!)\[([^\]]+)\]\(([^\)]+)\)', text)
        ]

    def _extract_images(self, text: str) -> list[dict]:
        return [
            {"alt": m.group(1) or "image", "url": m.group(2)}
            for m in re.finditer(r'!\[([^\]]*)\]\(([^\)]+)\)', text)
        ]

    def _extract_list_items(self, text: str) -> list[dict]:
        items = []
        for line_num, line in enumerate(text.splitlines(), start=1):
            m = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
            if not m:
                continue
            marker = m.group(2)
            items.append({
                "text": m.group(3).strip(),
                "level": len(m.group(1)) // 2,
                "ordered": marker[0].isdigit(),
                "line": line_num,
            })
        return items

    def _strip_code_blocks(self, text: str) -> str:
        return re.sub(r'```.*?```', '', text, flags=re.DOTALL).strip()

    def _build_chunks(self, sections: list[Section], code_blocks: list[CodeBlock]) -> list[dict]:
        chunks = []
        for sec in sections:
            if sec.content.strip():
                chunks.extend(_chunk(sec.content, chunk_type="section",
                                     meta={"section": sec.title, "level": sec.level}))
        for cb in code_blocks:
            chunks.append({"text": cb.code, "type": "code",
                           "language": cb.language, "start": 0, "end": len(cb.code)})
        return chunks


# Plain text files

class PlainTextParser:

    def parse(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")
        paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
        sections = [Section(title="", level=0, content=p) for p in paragraphs]
        chunks = []
        for p in paragraphs:
            chunks.extend(_chunk(p, chunk_type="paragraph"))

        words = len(text.split())
        return ParsedDocument(
            title        = path.stem.replace("-", " ").replace("_", " ").title(),
            file_path    = str(path),
            format       = "txt",
            raw_text     = text,
            sections     = sections,
            chunks       = chunks,
            metadata     = {
                "filename": path.name,
                "paragraph_count": len(paragraphs),
                "chunk_strategy": "sliding_window",
            },
            word_count   = words,
            reading_time_min = max(1, words // 250),
        )


# PDF files

class PDFParser:
    """
    Prefer pdfplumber so PDFs can contribute page chunks and table chunks.
    PyPDF2 stays as a fallback for environments where pdfplumber is missing.
    """

    def parse(self, path: Path) -> ParsedDocument:
        try:
            return self._parse_pdfplumber(path)
        except ImportError:
            log.warning("pdfplumber not installed — falling back to PyPDF2")
            return self._parse_pypdf2(path)

    def _parse_pdfplumber(self, path: Path) -> ParsedDocument:
        import pdfplumber

        all_text, tables, sections, chunks, headings = [], [], [], [], []

        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                all_text.append(page_text)
                headings.extend(self._extract_layout_headings(page, page_num))

                for tbl in page.extract_tables():
                    headers, rows = self._normalise_table(tbl)
                    if headers and rows:
                        tables.append(TableData(headers=headers, rows=rows,
                                                caption=f"Page {page_num}"))
                        table_text = " | ".join(headers) + "\n"
                        table_text += "\n".join(" | ".join(r) for r in rows[:20])
                        chunks.append({"text": table_text, "type": "table",
                                       "page": page_num, "start": 0, "end": len(table_text)})

                if page_text.strip():
                    sections.append(Section(title=f"Page {page_num}",
                                            level=0, content=page_text))
                    chunks.extend(_chunk(page_text, chunk_type="page",
                                         meta={"page": page_num}))

        full_text = "\n\n".join(all_text)
        words = len(full_text.split())

        return ParsedDocument(
            title        = path.stem.replace("-", " ").replace("_", " ").title(),
            file_path    = str(path),
            format       = "pdf",
            raw_text     = full_text,
            sections     = sections,
            chunks       = chunks,
            tables       = tables,
            metadata     = {
                "filename": path.name,
                "pages": len(all_text),
                "parser": "pdfplumber",
                "headings": headings,
                "table_count": len(tables),
            },
            word_count   = words,
            reading_time_min = max(1, words // 250),
        )

    def _normalise_table(self, table: Any) -> tuple[list[str], list[list[str]]]:
        """Clean pdfplumber table output before it becomes searchable text."""
        if not table or len(table) < 2:
            return [], []

        cleaned = [
            [str(cell or "").strip() for cell in row]
            for row in table
            if row and any(str(cell or "").strip() for cell in row)
        ]
        if len(cleaned) < 2:
            return [], []

        width = max(len(row) for row in cleaned)
        padded = [row + [""] * (width - len(row)) for row in cleaned]
        headers = [cell or f"Column {index + 1}" for index, cell in enumerate(padded[0])]
        rows = padded[1:]
        return headers, rows

    def _extract_layout_headings(self, page: Any, page_num: int) -> list[dict]:
        """Best-effort heading hints from pdfplumber font sizes."""
        try:
            words = page.extract_words(extra_attrs=["size"])
        except (TypeError, ValueError, AttributeError) as exc:
            log.debug("Could not extract PDF heading hints on page %s: %s", page_num, exc)
            return []
        sizes = [float(word.get("size", 0)) for word in words if word.get("text")]
        if not sizes:
            return []
        threshold = max(sum(sizes) / len(sizes) * 1.25, 13)
        headings = []
        for word in words:
            text = (word.get("text") or "").strip()
            size = float(word.get("size", 0) or 0)
            if text and size >= threshold and len(text) > 2:
                headings.append({"text": text, "page": page_num, "size": round(size, 2)})
        return headings[:20]

    def _parse_pypdf2(self, path: Path) -> ParsedDocument:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        pages = [p.extract_text() or "" for p in reader.pages]
        full_text = "\n\n".join(pages)
        sections = [Section(f"Page {i+1}", 0, t) for i, t in enumerate(pages) if t.strip()]
        chunks = []
        for i, t in enumerate(pages):
            chunks.extend(_chunk(t, chunk_type="page", meta={"page": i + 1}))

        words = len(full_text.split())
        return ParsedDocument(
            title        = path.stem.replace("-", " ").title(),
            file_path    = str(path),
            format       = "pdf",
            raw_text     = full_text,
            sections     = sections,
            chunks       = chunks,
            metadata     = {"filename": path.name, "pages": len(pages), "parser": "PyPDF2"},
            word_count   = words,
            reading_time_min = max(1, words // 250),
        )


# Word documents

class DocxParser:
    """
    python-docx plus direct DOCX XML reads for resumes and templates.

    Word templates often place visible content inside layout tables. Those
    layout tables should become searchable text, not data tables.
    """

    def parse(self, path: Path) -> ParsedDocument:
        from docx import Document

        doc = Document(str(path))

        sections, chunks, tables, images = [], [], [], []
        current_heading, current_level, buf = "Introduction", 0, []
        paragraph_count = 0

        for para in doc.paragraphs:
            style = para.style.name.lower()
            text  = para.text.strip()
            if not text:
                continue
            paragraph_count += 1

            if style.startswith("heading"):
                # flush previous section
                if buf:
                    content = "\n".join(buf)
                    sections.append(Section(current_heading, current_level, content))
                    chunks.extend(_chunk(content, chunk_type="section",
                                         meta={"section": current_heading}))
                    buf = []
                try:
                    current_level = int(style.split()[-1])
                except (ValueError, IndexError) as exc:
                    log.debug("Could not read DOCX heading level from style %r in %s: %s", style, path.name, exc)
                    current_level = 1
                current_heading = text
            else:
                buf.append(text)

        if buf:
            content = "\n".join(buf)
            sections.append(Section(current_heading, current_level, content))
            chunks.extend(_chunk(content, chunk_type="section"))

        for index, tbl in enumerate(doc.tables, 1):
            rows = self._normalise_table_rows(tbl)
            if self._is_data_table(rows):
                headers = rows[0]
                data    = rows[1:]
                tables.append(TableData(headers=headers, rows=data))
                table_text = " | ".join(headers) + "\n" + "\n".join(" | ".join(r) for r in data[:20])
                chunks.append({"text": table_text, "type": "table",
                               "start": 0, "end": len(table_text)})
            else:
                # Resume templates commonly use tables only for columns/layout.
                # Keep their text in the document, but do not report them as
                # extracted data tables in the UI.
                layout_text = self._table_text(rows)
                if layout_text:
                    sections.append(Section(f"Layout table {index}", 0, layout_text))
                    chunks.extend(_chunk(layout_text, chunk_type="section",
                                         meta={"section": f"Layout table {index}"}))

        for index, shape in enumerate(doc.inline_shapes, 1):
            images.append({
                "index": index,
                "type": str(getattr(shape, "type", "inline")),
            })

        heading_meta = [
            {"title": section.title, "level": section.level}
            for section in sections
            if section.level > 0
        ]
        app_props = self._read_app_properties(path)
        comments_count = self._count_comments(path)
        inherited_styles = self._count_inherited_styles(doc)
        full_text = "\n\n".join(s.content for s in sections)
        paragraph_count = self._count_visible_paragraphs(path) or paragraph_count
        xml_text = self._extract_visible_text(path)
        # Word's saved app.xml word count can be stale after edits. Count the
        # visible XML text so summaries match what the parser can actually use.
        extracted_words = self._count_words(xml_text or full_text)
        app_words = app_props.get("words", 0)

        return ParsedDocument(
            title        = sections[0].title if sections else path.stem,
            file_path    = str(path),
            format       = "docx",
            raw_text     = full_text,
            sections     = sections,
            chunks       = chunks,
            tables       = tables,
            images       = images,
            metadata     = {
                "filename": path.name,
                "headings": heading_meta,
                "paragraph_count": paragraph_count,
                "pages": app_props.get("pages", 0),
                "table_count": len(tables),
                "image_count": len(images),
                "comments_count": comments_count,
                "inherited_styles_count": inherited_styles,
                "extracted_word_count": extracted_words,
                "app_word_count": app_words,
                "word_count_source": "document_xml",
            },
            word_count   = extracted_words,
            reading_time_min = max(1, extracted_words // 250),
        )

    def _normalise_table_rows(self, table: Any) -> list[list[str]]:
        rows = []
        for row in table.rows:
            values = []
            for cell in row.cells:
                text = cell.text.strip()
                if values and values[-1] == text:
                    continue
                values.append(text)
            rows.append(values)
        return rows

    def _is_data_table(self, rows: list[list[str]]) -> bool:
        """Ignore empty/layout-only tables used to position text in templates."""
        if len(rows) < 2:
            return False
        width = max((len(row) for row in rows), default=0)
        if width < 2:
            return False
        cells = [cell for row in rows for cell in row if cell.strip()]
        non_empty_cells = len(cells)
        non_empty_rows = sum(1 for row in rows if any(cell.strip() for cell in row))
        if non_empty_cells < 4 or non_empty_rows < 2:
            return False
        word_counts = [len(cell.split()) for cell in cells]
        average_words = sum(word_counts) / len(word_counts)
        return average_words <= 14 and max(word_counts) <= 40

    def _table_text(self, rows: list[list[str]]) -> str:
        seen = set()
        pieces = []
        for row in rows:
            for cell in row:
                text = cell.strip()
                if text and text not in seen:
                    pieces.append(text)
                    seen.add(text)
        return "\n\n".join(pieces)

    def _read_app_properties(self, path: Path) -> dict[str, int]:
        """Read Word's saved page/word counts when the DOCX includes them."""
        try:
            with zipfile.ZipFile(path) as archive:
                data = archive.read("docProps/app.xml")
        except KeyError:
            log.debug("DOCX %s has no docProps/app.xml metadata", path.name)
            return {"pages": 0, "words": 0}
        except (OSError, zipfile.BadZipFile) as exc:
            log.warning("Could not read DOCX app properties from %s: %s", path.name, exc)
            return {"pages": 0, "words": 0}

        try:
            root = ElementTree.fromstring(data)
        except ElementTree.ParseError as exc:
            log.warning("Invalid DOCX app properties XML in %s: %s", path.name, exc)
            return {"pages": 0, "words": 0}

        values = {"pages": 0, "words": 0}
        for child in root:
            tag = child.tag.rsplit("}", 1)[-1].lower()
            if tag in values and child.text:
                try:
                    values[tag] = int(child.text)
                except ValueError as exc:
                    log.debug("Ignoring non-numeric DOCX %s value in %s: %s", tag, path.name, exc)
                    values[tag] = 0
        return values

    def _extract_visible_text(self, path: Path) -> str:
        """Read DOCX text nodes directly, including text missed by python-docx."""
        parts = [
            "word/document.xml",
            "word/footnotes.xml",
            "word/endnotes.xml",
            "word/comments.xml",
        ]
        try:
            with zipfile.ZipFile(path) as archive:
                parts.extend(name for name in archive.namelist() if re.match(r"word/(header|footer)\d+\.xml$", name))
                text = []
                for name in dict.fromkeys(parts):
                    try:
                        root = ElementTree.fromstring(archive.read(name))
                    except KeyError:
                        log.debug("DOCX %s has no %s part", path.name, name)
                        continue
                    except ElementTree.ParseError as exc:
                        log.warning("Skipping invalid DOCX XML part %s in %s: %s", name, path.name, exc)
                        continue
                    text.extend(node.text or "" for node in root.iter() if node.tag.endswith("}t"))
        except (OSError, zipfile.BadZipFile) as exc:
            log.warning("Could not read visible DOCX text from %s: %s", path.name, exc)
            return ""
        return " ".join(piece.strip() for piece in text if piece.strip())

    def _count_visible_paragraphs(self, path: Path) -> int:
        try:
            with zipfile.ZipFile(path) as archive:
                root = ElementTree.fromstring(archive.read("word/document.xml"))
        except KeyError:
            log.debug("DOCX %s has no word/document.xml part", path.name)
            return 0
        except ElementTree.ParseError as exc:
            log.warning("Invalid DOCX document XML in %s: %s", path.name, exc)
            return 0
        except (OSError, zipfile.BadZipFile) as exc:
            log.warning("Could not count DOCX paragraphs in %s: %s", path.name, exc)
            return 0

        count = 0
        for para in root.iter():
            if not para.tag.endswith("}p"):
                continue
            text = " ".join(node.text or "" for node in para.iter() if node.tag.endswith("}t")).strip()
            if text:
                count += 1
        return count

    def _count_words(self, text: str) -> int:
        return len(re.findall(r"\b[\w'-]+\b", text))

    def _count_comments(self, path: Path) -> int:
        try:
            with zipfile.ZipFile(path) as archive:
                data = archive.read("word/comments.xml")
        except KeyError:
            log.debug("DOCX %s has no comments.xml", path.name)
            return 0
        except (OSError, zipfile.BadZipFile) as exc:
            log.warning("Could not read DOCX comments from %s: %s", path.name, exc)
            return 0

        try:
            root = ElementTree.fromstring(data)
        except ElementTree.ParseError as exc:
            log.warning("Invalid DOCX comments XML in %s: %s", path.name, exc)
            return 0
        return sum(1 for node in root.iter() if node.tag.endswith("}comment"))

    def _count_inherited_styles(self, doc: Any) -> int:
        return sum(1 for style in doc.styles if getattr(style, "base_style", None) is not None)


# Python files

class PythonParser:
    """
    AST-based Python parser.
    Extracts: imports, functions, classes, docstrings, type hints.
    """

    def parse(self, path: Path) -> ParsedDocument:
        source = path.read_text(encoding="utf-8", errors="replace")

        imports, functions, classes, docstrings = [], [], [], []

        try:
            tree = ast.parse(source)
            for node in ast.walk(tree):

                # Imports
                if isinstance(node, ast.Import):
                    imports.extend(alias.name for alias in node.names)
                elif isinstance(node, ast.ImportFrom):
                    if node.module:
                        imports.append(node.module)

                # Functions
                elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                    args = [a.arg for a in node.args.args]
                    sig  = f"{node.name}({', '.join(args)})"
                    functions.append(sig)
                    doc  = ast.get_docstring(node)
                    if doc:
                        docstrings.append({"name": node.name, "doc": doc})

                # Classes
                elif isinstance(node, ast.ClassDef):
                    classes.append(node.name)
                    doc = ast.get_docstring(node)
                    if doc:
                        docstrings.append({"name": node.name, "doc": doc})

        except SyntaxError as e:
            log.warning("Python AST parse failed for %s: %s", path.name, e)

        # Build sections from top-level docstring + each class/function
        sections = []
        module_doc = ""
        try:
            tree = ast.parse(source)
            module_doc = ast.get_docstring(tree) or ""
        except SyntaxError as exc:
            log.debug("Skipping Python module docstring for invalid syntax in %s: %s", path.name, exc)

        if module_doc:
            sections.append(Section("Module", 1, module_doc))

        for d in docstrings:
            sections.append(Section(d["name"], 2, d["doc"]))

        # Chunks: source split by class/function boundaries + docstrings
        chunks = []
        if module_doc:
            chunks.extend(_chunk(module_doc, chunk_type="docstring", meta={"name": "module"}))
        for d in docstrings:
            chunks.extend(_chunk(d["doc"], chunk_type="docstring", meta={"name": d["name"]}))
        # Raw source as one big code chunk
        chunks.append({"text": source, "type": "code", "language": "python",
                       "start": 0, "end": len(source)})

        # Build readable summary text
        summary_parts = [f"Python module: {path.name}"]
        if imports:
            summary_parts.append(f"Imports: {', '.join(dict.fromkeys(imports))}")
        if classes:
            summary_parts.append(f"Classes: {', '.join(classes)}")
        if functions:
            summary_parts.append(f"Functions: {', '.join(f.split('(')[0] for f in functions[:20])}")
        if module_doc:
            summary_parts.append(module_doc)

        raw_text = "\n\n".join(summary_parts)
        words = len(raw_text.split())
        unique_imports = list(dict.fromkeys(imports))

        return ParsedDocument(
            title        = path.stem,
            file_path    = str(path),
            format       = "py",
            raw_text     = raw_text,
            sections     = sections,
            chunks       = chunks,
            code_blocks  = [CodeBlock("python", source)],
            imports      = unique_imports,
            functions    = functions,
            classes      = classes,
            entities     = _dependency_entities(unique_imports),
            metadata     = {
                "filename":      path.name,
                "import_count":  len(imports),
                "function_count":len(functions),
                "class_count":   len(classes),
                "dependencies":  _top_level_dependencies(unique_imports),
            },
            word_count   = words,
            reading_time_min = 1,
        )


# JavaScript and TypeScript files

class JavaScriptParser:
    """
    Regex-based JS/TS parser.
    Extracts: imports, exports, functions, classes, JSDoc comments.
    For production, consider using tree-sitter for full AST.
    """

    def parse(self, path: Path) -> ParsedDocument:
        source = path.read_text(encoding="utf-8", errors="replace")
        ext    = path.suffix.lower()

        # Imports
        import_patterns = [
            r'import\s+.*?\s+from\s+[\'"]([^\'"]+)[\'"]',
            r'require\s*\(\s*[\'"]([^\'"]+)[\'"]\s*\)',
            r'import\s*\(\s*[\'"]([^\'"]+)[\'"]\s*\)',
        ]
        imports = []
        for pat in import_patterns:
            imports.extend(re.findall(pat, source))

        # Functions
        func_patterns = [
            r'(?:function\s+(\w+)\s*\()',
            r'(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s+)?\(',
            r'(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s+)?function',
            r'(\w+)\s*\([^)]*\)\s*(?::\s*\w+\s*)?\{',   # method shorthand
        ]
        functions = []
        for pat in func_patterns:
            functions.extend(re.findall(pat, source))

        # Classes
        classes = re.findall(r'class\s+(\w+)', source)

        # JSDoc comments
        jsdoc_blocks = re.findall(r'/\*\*(.*?)\*/', source, re.DOTALL)
        docstrings = []
        for block in jsdoc_blocks:
            clean = re.sub(r'\s*\*\s?', ' ', block).strip()
            if len(clean) > 20:
                docstrings.append(clean)

        sections = []
        for i, doc in enumerate(docstrings):
            sections.append(Section(f"JSDoc {i+1}", 2, doc))

        summary_parts = [f"{'TypeScript' if ext == '.ts' else 'JavaScript'} module: {path.name}"]
        if imports:
            summary_parts.append(f"Imports: {', '.join(dict.fromkeys(imports))}")
        if classes:
            summary_parts.append(f"Classes: {', '.join(classes)}")
        if functions:
            unique_funcs = list(dict.fromkeys(f for f in functions if len(f) > 1))
            summary_parts.append(f"Functions: {', '.join(unique_funcs[:20])}")
        if docstrings:
            summary_parts.extend(docstrings[:5])

        raw_text = "\n\n".join(summary_parts)
        chunks = [{"text": raw_text, "type": "summary", "start": 0, "end": len(raw_text)}]
        chunks.append({"text": source, "type": "code", "language": ext.lstrip("."),
                       "start": 0, "end": len(source)})

        unique_imports = list(dict.fromkeys(imports))
        unique_functions = list(dict.fromkeys(f for f in functions if len(f) > 1))

        return ParsedDocument(
            title        = path.stem,
            file_path    = str(path),
            format       = ext.lstrip("."),
            raw_text     = raw_text,
            sections     = sections,
            chunks       = chunks,
            code_blocks  = [CodeBlock(ext.lstrip("."), source)],
            imports      = unique_imports,
            functions    = unique_functions,
            classes      = classes,
            entities     = _dependency_entities(unique_imports),
            metadata     = {
                "filename": path.name,
                "dependencies": _top_level_dependencies(unique_imports),
                "function_count": len(unique_functions),
                "class_count": len(classes),
            },
            word_count   = len(raw_text.split()),
            reading_time_min = 1,
        )


# JSON files

class JSONParser:
    """
    Extracts schema, keys, and values from JSON.
    Handles both objects and arrays.
    """

    def parse(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")

        try:
            data = json.loads(text)
        except json.JSONDecodeError as e:
            log.warning("JSON parse error in %s: %s", path.name, e)
            data = {}

        sections, chunks = [], []
        entities = self._extract_json_entities(data)

        if isinstance(data, dict):
            schema_desc = self._describe_object(data, depth=0)
            sections.append(Section("JSON Object", 1, schema_desc))
            chunks.extend(_chunk(schema_desc, chunk_type="schema"))

            # Flatten string values as searchable content
            str_values = self._extract_strings(data)
            if str_values:
                val_text = "\n".join(str_values[:100])
                sections.append(Section("Values", 2, val_text))
                chunks.extend(_chunk(val_text, chunk_type="values"))

        elif isinstance(data, list):
            desc = f"JSON array with {len(data)} items."
            if data:
                desc += f"\nSample: {json.dumps(data[0], ensure_ascii=False)[:500]}"
            sections.append(Section("JSON Array", 1, desc))
            chunks.extend(_chunk(desc, chunk_type="schema"))

        raw_text = "\n\n".join(s.content for s in sections) or text[:5000]

        return ParsedDocument(
            title        = path.stem,
            file_path    = str(path),
            format       = "json",
            raw_text     = raw_text,
            sections     = sections,
            chunks       = chunks,
            entities     = entities,
            metadata     = {"filename": path.name, "type": type(data).__name__},
            word_count   = len(raw_text.split()),
            reading_time_min = 1,
        )

    def _describe_object(self, obj: dict, depth: int = 0, prefix: str = "") -> str:
        lines = []
        for k, v in list(obj.items())[:50]:  # cap at 50 keys
            key_path = f"{prefix}.{k}" if prefix else k
            if isinstance(v, dict):
                lines.append(f"{key_path}: object ({len(v)} keys)")
                if depth < 2:
                    lines.append(self._describe_object(v, depth + 1, key_path))
            elif isinstance(v, list):
                lines.append(f"{key_path}: array[{len(v)}]")
            else:
                val_str = str(v)[:100]
                lines.append(f"{key_path}: {val_str}")
        return "\n".join(lines)

    def _extract_strings(self, obj, results=None) -> list[str]:
        if results is None:
            results = []
        if isinstance(obj, str) and len(obj) > 10:
            results.append(obj)
        elif isinstance(obj, dict):
            for v in obj.values():
                self._extract_strings(v, results)
        elif isinstance(obj, list):
            for item in obj[:100]:
                self._extract_strings(item, results)
        return results

    def _extract_json_entities(self, data: Any) -> list[dict]:
        """Pull package/API-style hints out of structured JSON."""
        if not isinstance(data, dict):
            return []

        entities = []
        for key in ("name", "package", "module"):
            value = data.get(key)
            if isinstance(value, str) and 2 <= len(value) <= 60:
                entities.append({"type": "library", "text": value})

        for key in ("dependencies", "devDependencies", "peerDependencies", "requires"):
            values = data.get(key)
            if isinstance(values, dict):
                for name in list(values.keys())[:25]:
                    entities.append({"type": "dependency", "text": name})
            elif isinstance(values, list):
                for name in values[:25]:
                    if isinstance(name, str):
                        entities.append({"type": "dependency", "text": name})

        for key in ("keywords", "tags", "classifiers"):
            values = data.get(key)
            if isinstance(values, list):
                for value in values[:20]:
                    if isinstance(value, str) and 2 <= len(value) <= 40:
                        entities.append({"type": "concept", "text": value})

        for key in ("summary", "description", "about"):
            value = data.get(key)
            if isinstance(value, str):
                for phrase in _known_concepts_in_text(value):
                    entities.append({"type": "concept", "text": phrase})

        return _dedupe_entity_dicts(entities)


# CSV files

class CSVParser:
    """
    Pandas-powered CSV parser.
    Extracts column names, data types, sample values, and basic stats.
    Falls back to stdlib csv if pandas unavailable.
    """

    def parse(self, path: Path) -> ParsedDocument:
        try:
            return self._parse_pandas(path)
        except ImportError:
            log.debug("pandas not installed; parsing CSV %s with stdlib csv", path.name)
            return self._parse_stdlib(path)

    def _parse_pandas(self, path: Path) -> ParsedDocument:
        import pandas as pd

        df = pd.read_csv(str(path), nrows=10000)  # cap for safety

        sections, chunks = [], []
        entities = self._extract_entities_from_frame(df)

        # Schema section
        schema_lines = [f"CSV file: {path.name}",
                        f"Rows: {len(df)}, Columns: {len(df.columns)}",
                        f"Columns: {', '.join(df.columns)}"]
        for col in df.columns:
            dtype = str(df[col].dtype)
            n_unique = df[col].nunique()
            sample = df[col].dropna().head(3).tolist()
            schema_lines.append(f"  {col} ({dtype}, {n_unique} unique): {sample}")

        schema_text = "\n".join(schema_lines)
        sections.append(Section("Schema", 1, schema_text))
        chunks.extend(_chunk(schema_text, chunk_type="schema"))

        # Numeric stats
        numeric_cols = df.select_dtypes(include="number")
        if not numeric_cols.empty:
            stats_text = str(numeric_cols.describe().round(2))
            sections.append(Section("Statistics", 1, stats_text))
            chunks.extend(_chunk(stats_text, chunk_type="statistics"))

        # Sample rows as text
        sample_text = df.head(10).to_string(index=False)
        sections.append(Section("Sample Data", 1, sample_text))
        chunks.extend(_chunk(sample_text, chunk_type="sample"))

        table = TableData(
            headers = list(df.columns),
            rows    = df.head(20).astype(str).values.tolist(),
        )

        raw_text = "\n\n".join(s.content for s in sections)

        return ParsedDocument(
            title        = path.stem.replace("_", " ").replace("-", " ").title(),
            file_path    = str(path),
            format       = "csv",
            raw_text     = raw_text,
            sections     = sections,
            chunks       = chunks,
            tables       = [table],
            entities     = entities,
            metadata     = {
                "filename":    path.name,
                "rows":        len(df),
                "columns":     list(df.columns),
                "column_count":len(df.columns),
                "entity_count": len(entities),
            },
            word_count   = len(raw_text.split()),
            reading_time_min = 1,
        )

    def _extract_entities_from_frame(self, df: Any) -> list[dict]:
        """Lightweight entity hints from low-cardinality text columns."""
        entities = []
        for col in df.columns:
            if not str(df[col].dtype).startswith("object"):
                continue
            values = [str(v).strip() for v in df[col].dropna().head(200).tolist()]
            candidates = [
                v for v in values
                if 2 <= len(v) <= 80 and not re.fullmatch(r"[-+]?\d+(\.\d+)?", v)
            ]
            for value in list(dict.fromkeys(candidates))[:20]:
                entities.append({"type": "csv_value", "text": value, "column": str(col)})
        return entities[:100]

    def _parse_stdlib(self, path: Path) -> ParsedDocument:
        with open(path, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows   = list(reader)

        if not rows:
            return ParsedDocument(title=path.stem, file_path=str(path), format="csv",
                                  raw_text="", metadata={"filename": path.name})

        headers = rows[0]
        data    = rows[1:21]
        table   = TableData(headers=headers, rows=data)

        text = f"CSV: {path.name}\nColumns: {', '.join(headers)}\n"
        text += "\n".join(", ".join(r) for r in data[:10])

        return ParsedDocument(
            title    = path.stem,
            file_path= str(path),
            format   = "csv",
            raw_text = text,
            sections = [Section("Sample Data", 1, text)],
            chunks   = _chunk(text, chunk_type="csv"),
            tables   = [table],
            metadata = {
                "filename": path.name,
                "columns": headers,
                "rows": max(0, len(rows) - 1),
                "column_count": len(headers),
            },
        )


# HTML files

class HTMLParser:
    """
    BeautifulSoup: strips navigation, ads, footers — extracts article body.
    """

    SKIP_TAGS = {"script", "style", "nav", "footer", "header",
                 "aside", "advertisement", "form", "noscript"}

    def parse(self, path: Path) -> ParsedDocument:
        from bs4 import BeautifulSoup

        html = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")

        # Extract title
        title = ""
        if soup.title:
            title = soup.title.string or ""
        if not title:
            h1 = soup.find("h1")
            title = h1.get_text(strip=True) if h1 else path.stem

        # Remove noisy tags
        for tag in soup(list(self.SKIP_TAGS)):
            tag.decompose()

        # Try to find main article content
        main = (soup.find("article")
                or soup.find("main")
                or soup.find(id=re.compile(r"content|main|article", re.I))
                or soup.find("body")
                or soup)

        # Extract headings as sections
        sections, links = [], []
        for tag in main.find_all(["h1", "h2", "h3", "h4", "p", "li", "a"]):
            if tag.name in ("h1", "h2", "h3", "h4"):
                level = int(tag.name[1])
                sections.append(Section(tag.get_text(strip=True), level, ""))
            elif tag.name == "a" and tag.get("href"):
                links.append({"text": tag.get_text(strip=True), "url": tag["href"]})

        # Full text
        text = " ".join(main.get_text(separator=" ").split())

        # Attach text to last section
        if sections:
            sections[-1].content = text
        else:
            sections = [Section(title, 1, text)]

        chunks = _chunk(text, chunk_type="paragraph")
        words  = len(text.split())

        return ParsedDocument(
            title        = title.strip(),
            file_path    = str(path),
            format       = "html",
            raw_text     = text,
            sections     = sections,
            chunks       = chunks,
            links        = links,
            metadata     = {"filename": path.name},
            word_count   = words,
            reading_time_min = max(1, words // 250),
        )


# Parser lookup

PARSERS = {
    ".md":   MarkdownParser(),
    ".txt":  PlainTextParser(),
    ".pdf":  PDFParser(),
    ".docx": DocxParser(),
    ".py":   PythonParser(),
    ".js":   JavaScriptParser(),
    ".ts":   JavaScriptParser(),
    ".json": JSONParser(),
    ".csv":  CSVParser(),
    ".html": HTMLParser(),
    ".htm":  HTMLParser(),
}


def parse_document(file_path: str) -> ParsedDocument:
    """
    Parse any supported document type.

    Auto-selects the correct parser based on file extension.
    Returns a unified ParsedDocument regardless of input format.

    Args:
        file_path: absolute or relative path to the file

    Returns:
        ParsedDocument with raw_text, sections, chunks, and metadata

    Raises:
        ValueError: if the file extension is not supported
        FileNotFoundError: if the file does not exist
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    parser = PARSERS.get(ext)

    if not parser:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(sorted(PARSERS))}"
        )

    log.info("Parsing %s as %s", path.name, ext)
    doc = parser.parse(path)

    log.info(
        "Parsed %s: %d words, %d sections, %d chunks",
        path.name, doc.word_count, len(doc.sections), len(doc.chunks),
    )
    return doc


def _top_level_dependencies(imports: list[str]) -> list[str]:
    """Normalize imports into dependency names useful for graph nodes."""
    dependencies = []
    for item in imports:
        root = _normalise_dependency_name(item)
        if root:
            dependencies.append(root)
    return list(dict.fromkeys(dependencies))


def _normalise_dependency_name(value: str) -> str:
    item = value.strip().split()[0].strip("'\"")
    if not item or item.startswith((".", "/")):
        return ""
    if item.startswith("@"):
        parts = item.split("/")
        return "/".join(parts[:2]) if len(parts) >= 2 else item
    return item.split("/")[0].split(".")[0]


def _dependency_entities(imports: list[str]) -> list[dict]:
    return [
        {"type": "dependency", "text": name}
        for name in _top_level_dependencies(imports)
    ]


def _known_concepts_in_text(text: str) -> list[str]:
    """Extract only clear technical concepts from noisy metadata prose."""
    concepts = [
        "API",
        "REST",
        "HTTP",
        "JSON",
        "GraphQL",
        "database",
        "semantic search",
        "knowledge graph",
        "machine learning",
        "data structure",
    ]
    lower = text.lower()
    return [concept for concept in concepts if concept.lower() in lower]


def _dedupe_entity_dicts(entities: list[dict]) -> list[dict]:
    seen = set()
    result = []
    for entity in entities:
        key = (str(entity.get("type", "")).lower(), str(entity.get("text", "")).lower())
        if key in seen or not key[1]:
            continue
        seen.add(key)
        result.append(entity)
    return result


def parsed_document_to_legacy_dict(doc: ParsedDocument) -> dict[str, Any]:
    """Adapter for older code paths that still expect dict parser output."""
    extra = {
        "sections": [asdict(section) for section in doc.sections],
        "code_blocks": [asdict(block) for block in doc.code_blocks],
        "tables": [asdict(table) for table in doc.tables],
        "links": doc.links,
        "images": doc.images,
        "list_items": doc.list_items,
        "imports": doc.imports,
        "functions": doc.functions,
        "classes": doc.classes,
        "entities": doc.entities,
    }
    return {
        "content": doc.raw_text,
        "chunks": doc.chunks,
        "metadata": {
            **doc.metadata,
            "filename": Path(doc.file_path).name,
            "format": doc.format,
            "size": Path(doc.file_path).stat().st_size,
            "num_chunks": len(doc.chunks),
            "word_count": doc.word_count,
            "reading_time_min": doc.reading_time_min,
            "title": doc.title,
        },
        "extra": extra,
    }


class DocumentParser:
    """Compatibility wrapper around the unified parser."""

    def parse(self, file_path: str) -> dict[str, Any]:
        return parsed_document_to_legacy_dict(parse_document(file_path))
